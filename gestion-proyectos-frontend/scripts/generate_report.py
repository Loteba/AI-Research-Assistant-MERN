from docx import Document
from docx.shared import Pt

report_path = 'reports/test_coverage_report.docx'

doc = Document()

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)

add_heading('Informe de Pruebas y Cobertura', level=1)
add_paragraph('Fecha: 2025-10-06')
add_paragraph('Proyecto: gestion-proyectos-frontend')

add_heading('Resumen ejecutivo', level=2)
add_paragraph('Se actualizó la suite de pruebas unitarias e integración para alcanzar y superar la cobertura requerida. Se añadieron tests, se corrigieron condiciones de carrera y se mejoró la robustez de pruebas asíncronas.', False)

add_heading('Detalles de lo cubierto', level=2)
add_heading('1. Cobertura superior al 70%', level=3)
add_paragraph('Estado: Done')
add_paragraph('Evidencia: Tras añadir tests, la cobertura global quedó:')

# Table with coverage numbers
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Métrica'
hdr_cells[1].text = 'Valor'
rows = [
    ('Statements', '80.03%'),
    ('Branches', '70.4%'),
    ('Functions', '73.42%'),
    ('Lines', '80.5%'),
]
for name, val in rows:
    r = table.add_row().cells
    r[0].text = name
    r[1].text = val

add_heading('2. Pruebas de interacción complejas', level=3)
add_paragraph('Estado: Good (Partial)')
add_paragraph('Se mantiene un test de integración principal: `src/pages/__tests__/LibraryPage.integration.test.js`. Cubre flujo de subida de archivo, debounce y refetch. Se hizo la prueba más robusta para aceptar múltiples llamadas a getItems y evitar condiciones de carrera.')

add_heading('3. Mocking de APIs y Contextos', level=3)
add_paragraph('Estado: Mostly done / Partial')
add_paragraph('Se emplearon mocks por test para servicios clave:')
add_paragraph('- `jest.mock` para `libraryService` en tests de integración y unitarios.')
add_paragraph('- Se añadió test para `AuthContext` que espía `useNavigate` en tiempo de ejecución y verifica login/logout y localStorage.')
add_paragraph('Recomendación: Añadir mock global centralizado para `apiClient` en `src/setupTests.js` si se desean tests de integración que aseguren aislamiento completo.')

add_heading('4. Pruebas de accesibilidad (a11y)', level=3)
add_paragraph('Estado: Not covered / Partial')
add_paragraph('Observación: `jest-axe` está en `devDependencies` pero no se añadieron tests de a11y durante esta iteración. Recomendación: añadir 1-2 tests con `jest-axe` (por ejemplo para `LibraryPage` y `SummaryCard`) como plantilla para el resto.')

add_heading('Cambios realizados (archivos clave)', level=2)
files = [
    'src/pages/__tests__/LibraryPage.integration.test.js  (ajustes async y mocks)',
    'src/context/__tests__/AuthContext.test.js  (login/logout, localStorage, navigate spy)',
    'src/components/ai/__tests__/SummaryCard.test.jsx  (parsing, copy, download, regenerate)',
]
for f in files:
    add_paragraph('- ' + f)

add_heading('Warnings y observaciones de test run', level=2)
add_paragraph('- Aparición de warnings `act(...)` en algunos tests que hacen updates asíncronos (p.ej. UploadItemModal, SummaryCard). No rompen la suite pero conviene envolver actualizaciones en `act` o await si se quiere limpiar la salida.')
add_paragraph('- Mensajes de `console.error` esperados en tests que validan manejo de errores (se simularon fallos en `libraryService.getItems`).')

add_heading('Comandos útiles', level=2)
add_paragraph('Ejecutar tests con cobertura (Windows PowerShell):')
add_paragraph("cd 'd:\\Proyectos\\AI\\AI-Research-Assistant-MERN-main\\gestion-proyectos-frontend'\nnpm run test:cov")

add_heading('Recomendaciones y próximos pasos', level=2)
add_paragraph('1. Añadir tests de accesibilidad con `jest-axe` para componentes/páginas clave.')
add_paragraph('2. Centralizar mocks de red (por ejemplo mockear `apiClient` en `src/setupTests.js`).')
add_paragraph('3. Añadir 1 E2E (Cypress/Playwright) para flujos críticos si se desea mayor confianza end-to-end.')

# Save document
import os
os.makedirs(os.path.dirname(report_path), exist_ok=True)
doc.save(report_path)
print('Report generated at', report_path)
